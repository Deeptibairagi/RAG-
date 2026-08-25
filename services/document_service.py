import os
import tempfile
import requests


# ==========================================================
# PDF
# ==========================================================

def extract_pdf(uploaded_file):

    from pypdf import PdfReader

    reader = PdfReader(uploaded_file)

    pages = []

    for page in reader.pages:

        text = page.extract_text()

        if text:

            pages.append(text)

    return "\n\n".join(pages)


# ==========================================================
# DOCX
# ==========================================================

def extract_docx(uploaded_file):

    from docx import Document

    document = Document(uploaded_file)

    paragraphs = []

    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            paragraphs.append(paragraph.text)

    return "\n".join(paragraphs)


# ==========================================================
# TXT
# ==========================================================

def extract_txt(uploaded_file):

    return uploaded_file.getvalue().decode("utf-8", errors="ignore")


# ==========================================================
# UPLOAD
# ==========================================================

def process_uploaded_file(uploaded_file):

    filename = uploaded_file.name.lower()

    try:

        if filename.endswith(".pdf"):

            return extract_pdf(uploaded_file)

        if filename.endswith(".docx"):

            return extract_docx(uploaded_file)

        if filename.endswith(".txt"):

            return extract_txt(uploaded_file)

        return ("Unsupported file format.")

    except Exception as e:

        return (f"Could not process file: {e}")


# ==========================================================
# URL
# ==========================================================

def load_http_file(url):

    temp_path = None

    try:

        response = requests.get(url, timeout=20 )

        response.raise_for_status()

        content_type = response.headers.get("Content-Type", "").lower()

        data = response.content

        # --------------------------------------------------
        # PDF
        # --------------------------------------------------

        if ("pdf" in content_type or url.lower().endswith(".pdf")):

            from pypdf import PdfReader

            temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

            temp.write(data)
            temp.close()

            temp_path = temp.name

            reader = PdfReader(temp_path)

            pages = []

            for page in reader.pages:

                text = page.extract_text()

                if text:

                    pages.append(text)

            return "\n\n".join(pages)

        # --------------------------------------------------
        # DOCX
        # --------------------------------------------------

        if ("word" in content_type or url.lower().endswith(".docx")):

            from docx import Document

            temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")

            temp.write(data)
            temp.close()

            temp_path = temp.name

            document = Document(temp_path)

            paragraphs = []

            for paragraph in document.paragraphs:

                if paragraph.text.strip():

                    paragraphs.append(paragraph.text)

            return "\n".join(paragraphs)

        return data.decode("utf-8", errors="ignore")

    except Exception as e:

        return (f"Could not load URL: {e}")

    finally:

        if (temp_path and os.path.exists(temp_path)):

            os.unlink(temp_path)


# ==========================================================
# BUILD CONTEXT
# ==========================================================

def build_document_context(document_text, source_name, source_type):

    if source_type == "file":

        return (
            "\n\n"
            "UPLOADED DOCUMENT:\n"
            f"File: {source_name}\n"
            "--------------------------------\n"
            f"{document_text}\n"
            "--------------------------------\n"
        )

    return (
        "\n\n"
        "UPLOADED DOCUMENT FROM URL:\n"
        f"URL: {source_name}\n"
        "--------------------------------\n"
        f"{document_text}\n"
        "--------------------------------\n"
    )